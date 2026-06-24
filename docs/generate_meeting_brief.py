#!/opt/homebrew/Cellar/python@3.11/3.11.14_2/Frameworks/Python.framework/Versions/3.11/bin/python3.11
"""
Shazlo — App Store Pre-Deployment Meeting Brief
Generates: docs/Shazlo_Meeting_Brief.docx
Run: python3.11 docs/generate_meeting_brief.py
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colours ─────────────────────────────────────────────────────────────────
C_BLACK      = RGBColor(0x0D, 0x0D, 0x0D)
C_GOLD       = RGBColor(0xEE, 0xBA, 0x2B)
C_DARK       = RGBColor(0x1E, 0x1E, 0x1E)
C_BODY       = RGBColor(0x3A, 0x3A, 0x3A)
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_RED        = RGBColor(0xC6, 0x28, 0x28)
C_ORANGE     = RGBColor(0xE6, 0x5C, 0x00)
C_BLUE       = RGBColor(0x15, 0x65, 0xC0)
C_GREEN      = RGBColor(0x2E, 0x7D, 0x32)
C_LIGHT_GREY = RGBColor(0xF7, 0xF7, 0xF7)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_str)
    tcPr.append(shd)


def run(para, text, bold=False, italic=False, size=10, color=C_BODY, font='Calibri'):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(16)
    r.font.color.rgb = C_DARK
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    # gold underline rule
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'EEBA2B')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(doc, text, color=C_DARK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    return p


def body(doc, text, size=10, color=C_BODY, after=5):
    p = doc.add_paragraph()
    run(p, text, size=size, color=color)
    p.paragraph_format.space_after = Pt(after)
    return p


def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    run(p, text, size=size, color=C_BODY)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(1)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p


def table(doc, headers, rows, hdr_bg='1E1E1E', col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hcells = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hcells[i], hdr_bg)
        p = hcells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = p.add_run(h)
        rn.bold = True
        rn.font.name = 'Calibri'
        rn.font.size = Pt(9)
        rn.font.color.rgb = C_WHITE
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        bg = 'F7F7F7' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            rn = p.add_run(str(val))
            rn.font.name = 'Calibri'
            rn.font.size = Pt(9)
            rn.font.color.rgb = C_DARK
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ─── Sections ────────────────────────────────────────────────────────────────

def cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SHAZLO')
    r.bold = True; r.font.size = Pt(36); r.font.color.rgb = C_GOLD; r.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('App Store Pre-Deployment — Meeting Brief')
    r.font.size = Pt(16); r.font.color.rgb = C_DARK; r.font.name = 'Calibri'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('──────────────────────────────────────')
    r.font.size = Pt(11); r.font.color.rgb = C_GOLD

    doc.add_paragraph()
    for label, value in [
        ('Date',     str(date.today())),
        ('App',      'Shazlo — AI Fashion Discovery (React Native 0.80.0)'),
        ('Platform', 'iOS App Store'),
        ('Status',   '🔴  Blocking issues must be resolved before submission'),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, f'{label}:  ', bold=True, size=11, color=C_DARK)
        run(p, value, size=11, color=C_BODY)

    doc.add_page_break()


def sec_overview(doc):
    h1(doc, 'Overview')
    body(doc,
         'The Shazlo iOS build is architecturally sound and the Android version is already live. '
         'Before submitting to the App Store there are four issues that Apple will reject the app for, '
         'plus two code-level bugs that will cause crashes on iOS. '
         'This document covers only those items — everything else can be addressed post-launch.',
         size=10, after=8)

    table(doc,
        ['Area', 'Blocking for App Store?', 'Severity'],
        [
            ['Sign in with Apple missing',           'Yes — Guideline 4.8',    '🔴 Guaranteed Rejection'],
            ['Google Sign-In not wired up',          'Yes — tied to above',    '🔴 Rejection / Non-functional'],
            ['Hook called inside useEffect',         'Yes — iOS crash',         '🔴 App Crash'],
            ['QueryClient inside App component',     'No — but UX breakage',   '🟠 High'],
            ['Hardcoded fake product details',       'Possible — Guideline 2.3','🟠 High'],
            ['console.log on every swipe render',   'No — performance',        '🟡 Medium'],
        ],
        col_widths=[2.6, 2.1, 2.0],
    )


def sec_blockers(doc):
    h1(doc, 'Blocking Issues')

    # ── 1. Sign in with Apple ──────────────────────────────────────────────
    h2(doc, '1  Sign in with Apple — Guaranteed Rejection')
    body(doc,
         'Apple Guideline 4.8: any app that offers a third-party login option (Google, Facebook, Twitter) '
         'must also offer Sign in with Apple as an equivalent option. '
         'The Google Sign-In SDK (@react-native-google-signin) is installed and referenced. '
         'Apple\'s automated checker flags this before a human reviewer even sees the app.',
         after=6)

    p = doc.add_paragraph()
    run(p, 'What needs to happen:', bold=True, size=10, color=C_DARK)
    bullet(doc, 'Install @invertase/react-native-apple-authentication')
    bullet(doc, 'Enable "Sign In with Apple" capability in Xcode → Signing & Capabilities')
    bullet(doc, 'Add an Apple Sign-In button alongside the Google button on the Onboard screen')
    bullet(doc, 'Backend needs a /v1/auth/apple-login endpoint that verifies Apple\'s identity token')
    bullet(doc, 'Must be tested on a real iPhone (Apple auth will not work on Simulator)')
    body(doc, 'Estimated effort: 1 day (frontend + backend combined)', size=9, color=C_BLUE, after=10)

    # ── 2. Google Sign-In not wired up ────────────────────────────────────
    h2(doc, '2  Google Sign-In Is Incomplete')
    body(doc,
         'The Google Sign-In button is currently commented out in Onboard.jsx. '
         'The handler (onGoogleButtonPress) successfully obtains an idToken from Google but '
         'never sends it to the backend and never dispatches the login action to Redux. '
         'As a result, tapping the Google button does nothing useful.',
         after=6)

    p = doc.add_paragraph()
    run(p, 'Current code (Onboard.jsx lines 156–172):', bold=True, size=10, color=C_DARK)
    for line in [
        'const onGoogleButtonPress = async () => {',
        '  const userInfo = await GoogleSignin.signIn();',
        '  const idToken = userInfo.data?.idToken;',
        '  // idToken obtained — but nothing sent to backend.',
        '  // setlogin never dispatched. User stays logged out.',
        '};',
        '',
        '// Button is commented out — never rendered:',
        '{/* <GoogleButton onPress={onGoogleButtonPress}/> */}',
    ]:
        code(doc, line)

    body(doc, '', after=4)
    p = doc.add_paragraph()
    run(p, 'What needs to happen:', bold=True, size=10, color=C_DARK)
    bullet(doc, 'POST the idToken to /v1/auth/google-login after obtaining it')
    bullet(doc, 'Dispatch setlogin({ user: data.user }) from the response')
    bullet(doc, 'Navigate to Home screen after successful login')
    bullet(doc, 'Uncomment the <GoogleButton> in the JSX')
    body(doc, 'Estimated effort: 2–3 hours', size=9, color=C_BLUE, after=10)

    # ── 3. useSelector inside useEffect ───────────────────────────────────
    h2(doc, '3  useSelector Called Inside useEffect — iOS Crash')
    body(doc,
         'In Home.tsx (line 162), useSelector is called inside a useEffect callback. '
         'This violates React\'s Rules of Hooks. On Android the error is silently swallowed '
         'in most debug builds, which is why the app works fine there. '
         'On iOS production builds React enforces this strictly and throws '
         '"Invalid hook call" — crashing the Home screen entirely.',
         after=6)

    p = doc.add_paragraph()
    run(p, 'Broken code (Home.tsx line 159–173):', bold=True, size=10, color=C_DARK)
    for line in [
        'useEffect(() => {',
        '  const loadGender = async () => {',
        '    const user = useSelector((state:any)=>state.auth.user);  // ❌ CRASH',
        '    if (user.preferred_gender) setGender(user.preferred_gender);',
        '  };',
        '  loadGender();',
        '}, []);',
    ]:
        code(doc, line)

    body(doc, '', after=4)
    p = doc.add_paragraph()
    run(p, 'The fix (2 lines):', bold=True, size=10, color=C_DARK)
    for line in [
        '// user is already selected at line 28 — just use that reference:',
        'const user = useSelector((state:any)=>state.auth.user);  // ✅ top level',
        '',
        'useEffect(() => {',
        '  if (user?.preferred_gender) setGender(user.preferred_gender);',
        '}, [user?.preferred_gender]);',
    ]:
        code(doc, line)
    body(doc, 'Estimated effort: 15 minutes', size=9, color=C_BLUE, after=10)


def sec_high(doc):
    h1(doc, 'High Priority (Fix Before or Shortly After Launch)')

    # ── QueryClient ──────────────────────────────────────────────────────
    h2(doc, '4  QueryClient Recreated on Every Re-render')
    body(doc,
         'In App.jsx line 136, new QueryClient() is called inside the App component body. '
         'This creates a brand-new QueryClient instance every time App re-renders, '
         'which destroys all cached data (cart contents, product lists). '
         'Users will see their cart appear to empty randomly.',
         after=6)

    p = doc.add_paragraph()
    run(p, 'Fix:', bold=True, size=10, color=C_DARK)
    for line in [
        '// Move outside the App function (module level):',
        'const queryClient = new QueryClient();   // ✅',
        '',
        'const App = () => {',
        '  return <QueryClientProvider client={queryClient}>...</QueryClientProvider>;',
        '};',
    ]:
        code(doc, line)
    body(doc, 'Estimated effort: 5 minutes', size=9, color=C_BLUE, after=10)

    # ── Hardcoded product info ────────────────────────────────────────────
    h2(doc, '5  Every Product Shows Hardcoded Fake Details')
    body(doc,
         'In SwipeUI.jsx (lines 1688–1706), the back-of-card information for every single product '
         'displays the same hardcoded text: "100% premium cotton", "Machine wash cold", '
         '"Easy 7-day return policy" — regardless of the actual product. '
         'Apple reviewers browse the app and can flag this under Guideline 2.3 (Accurate Metadata / '
         'misleading content). More importantly it is false information shown to real users.',
         after=6)

    p = doc.add_paragraph()
    run(p, 'What needs to happen:', bold=True, size=10, color=C_DARK)
    bullet(doc, 'If material / care / return data exists in the database, display it from the API response')
    bullet(doc, 'If that data does not exist yet, hide the section entirely rather than showing placeholder text')
    bullet(doc, 'Never display hardcoded product attributes for real products')
    body(doc, 'Estimated effort: 2–4 hours (depends on whether backend fields exist)', size=9, color=C_BLUE, after=10)


def sec_medium(doc):
    h1(doc, 'Medium Priority (Clean Up for Production Build)')

    h2(doc, '6  console.log Runs on Every Swipe')
    body(doc,
         'SwipeUI.jsx logs the entire items array (console.log(items) at line 1054) and '
         'console.log(brand) at line 73 on every render. '
         'Home.tsx also logs user.preferred_gender on every effect run. '
         'This has no App Store rejection risk but does degrade performance on older iPhones '
         'and exposes product inventory data in device logs.',
         after=6)

    p = doc.add_paragraph()
    run(p, 'Fix — strip all logs in production builds via babel.config.js:', bold=True, size=10, color=C_DARK)
    for line in [
        '// babel.config.js',
        'module.exports = {',
        '  presets: [\'module:@react-native/babel-preset\'],',
        '  env: {',
        '    production: { plugins: [\'transform-remove-console\'] },',
        '  },',
        '};',
        '',
        '// Install:  npm install --save-dev babel-plugin-transform-remove-console',
    ]:
        code(doc, line)
    body(doc, 'Estimated effort: 10 minutes', size=9, color=C_BLUE, after=8)

    h2(doc, '7  Duplicate Route Key in Deep Linking Config (App.jsx line 158–160)')
    body(doc,
         'The deep linking config declares the Liked screen key twice. '
         'The second declaration overwrites the first, so one of the two routes silently stops working.',
         after=6)
    for line in [
        '// App.jsx  — current (broken):',
        'Liked: \'liked\',      // line 158 — overwritten, never reachable',
        'Liked: \'product/:id\', // line 160 — this wins',
        '',
        '// Fix — use distinct keys:',
        'Liked:   \'liked\',',
        'Product: \'product/:id\',',
    ]:
        code(doc, line)
    body(doc, 'Estimated effort: 5 minutes', size=9, color=C_BLUE, after=4)


def sec_summary(doc):
    doc.add_page_break()
    h1(doc, 'Summary & Effort Estimate')

    table(doc,
        ['#', 'Issue', 'Blocks App Store?', 'Est. Effort'],
        [
            ['1', 'Sign in with Apple (Guideline 4.8)',          'Yes — rejection',          '1 day'],
            ['2', 'Google Sign-In not wired to backend/Redux',   'Yes — tied to #1',         '2–3 hrs'],
            ['3', 'useSelector inside useEffect (iOS crash)',    'Yes — crash on Home screen','15 min'],
            ['4', 'QueryClient inside App component',            'No — cart data loss',       '5 min'],
            ['5', 'Hardcoded fake product details',              'Possibly — Guideline 2.3',  '2–4 hrs'],
            ['6', 'console.log on every render',                 'No — performance issue',    '10 min'],
            ['7', 'Duplicate deep-link route key',               'No — broken route',         '5 min'],
        ],
        col_widths=[0.3, 2.9, 1.9, 1.1],
    )

    body(doc,
         'Total estimated effort for items 1–3 (true blockers): 1–2 engineering days. '
         'Items 4–7 can be completed in under an hour alongside any other cleanup.',
         size=10, after=10)

    h2(doc, 'Recommended Order of Work')
    for i, step in enumerate([
        'Fix BUG-003 first (useSelector) — 15 minutes, highest crash risk',
        'Complete Google Sign-In end-to-end wiring — 2–3 hours',
        'Implement Sign in with Apple (frontend + backend) — 1 day',
        'Fix QueryClient location — 5 minutes',
        'Remove/replace hardcoded product details — 2–4 hours',
        'Add babel-plugin-transform-remove-console — 10 minutes',
        'Fix duplicate deep-link key — 5 minutes',
    ], 1):
        bullet(doc, f'{i}.  {step}')


def footer(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, '— End of Brief —', bold=True, size=9.5, color=C_GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, f'Prepared {date.today()}  ·  Shazlo Engineering', size=9, color=RGBColor(0xAA, 0xAA, 0xAA))


# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    cover(doc)
    sec_overview(doc)
    sec_blockers(doc)
    sec_high(doc)
    sec_medium(doc)
    sec_summary(doc)
    footer(doc)

    out = os.path.join(os.path.dirname(__file__), 'Shazlo_Meeting_Brief.docx')
    doc.save(out)
    print(f'✅  Saved: {out}')
    print('    Open in Microsoft Word or Google Docs (File → Import)')


if __name__ == '__main__':
    main()
