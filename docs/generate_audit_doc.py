#!/opt/homebrew/Cellar/python@3.11/3.11.14_2/Frameworks/Python.framework/Versions/3.11/bin/python3.11
"""
Shazlo iOS App Store — Production Engineering Audit Document Generator
Generates: docs/Shazlo_Engineering_Audit.docx
Opens in: Microsoft Word, Google Docs (File → Import), LibreOffice

Run:
  python3.11 docs/generate_audit_doc.py
  OR if python-docx is not on system python:
  /opt/homebrew/Cellar/python@3.11/3.11.14_2/Frameworks/Python.framework/Versions/3.11/bin/python3.11 docs/generate_audit_doc.py
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Colour palette ──────────────────────────────────────────────────────────
C_BLACK       = RGBColor(0x00, 0x00, 0x00)
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_GOLD        = RGBColor(0xEE, 0xBA, 0x2B)   # Shazlo brand
C_DARK_GREY   = RGBColor(0x1E, 0x1E, 0x1E)
C_MID_GREY    = RGBColor(0x4A, 0x4A, 0x4A)
C_LIGHT_GREY  = RGBColor(0xF5, 0xF5, 0xF5)
C_RED         = RGBColor(0xD3, 0x2F, 0x2F)
C_ORANGE      = RGBColor(0xF5, 0x7C, 0x00)
C_YELLOW      = RGBColor(0xF9, 0xA8, 0x25)
C_GREEN       = RGBColor(0x2E, 0x7D, 0x32)
C_BLUE        = RGBColor(0x15, 0x65, 0xC0)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via XML shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_run(para, text, bold=False, italic=False, size=None, color=None, font='Calibri'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def heading(doc, text, level=1, color=C_DARK_GREY, size=None, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    sizes = {1: 20, 2: 16, 3: 13, 4: 11}
    add_run(p, text, bold=bold, size=size or sizes.get(level, 11), color=color)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(doc, text, size=10, color=C_MID_GREY, space_after=6):
    p = doc.add_paragraph()
    add_run(p, text, size=size, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, level=0, size=10):
    p = doc.add_paragraph(style='List Bullet')
    add_run(p, text, size=size, color=C_MID_GREY)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p


def code_block(doc, text, size=8.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x18, 0x18, 0x18)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    # Light grey background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p


def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


def make_table(doc, headers, rows, header_bg='1E1E1E', header_color=C_WHITE,
               col_widths=None, stripe_bg='F5F5F5'):
    """Build a styled table. headers=list[str], rows=list[list[str]]."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], header_bg)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=9, color=header_color)

    # Data rows
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            if ri % 2 == 0:
                set_cell_bg(row_cells[ci], stripe_bg)
            else:
                set_cell_bg(row_cells[ci], 'FFFFFF')
            p = row_cells[ci].paragraphs[0]
            add_run(p, str(val), size=9, color=C_DARK_GREY)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def severity_badge(text):
    badges = {
        'CRITICAL': '🔴 CRITICAL',
        'HIGH':     '🟠 HIGH',
        'MEDIUM':   '🟡 MEDIUM',
        'LOW':      '🟢 LOW',
        'INFO':     '🔵 INFO',
    }
    return badges.get(text.upper(), text)


# ─── Document sections ───────────────────────────────────────────────────────

def cover_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'SHAZLO', bold=True, size=42, color=C_GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'iOS App Store', bold=False, size=24, color=C_DARK_GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'Production Engineering Audit', bold=True, size=24, color=C_DARK_GREY)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━', size=12, color=C_GOLD)

    doc.add_paragraph()

    meta = [
        ('Document Type',   'Production Readiness & Deployment Audit'),
        ('App',             'Shazlo — AI Fashion Discovery'),
        ('Platform',        'iOS (React Native 0.80.0 / New Architecture)'),
        ('Backend',         'Django 5.2.3  —  https://api.shazlo.store/'),
        ('Bundle ID',       'com.shazlo.mobile'),
        ('iOS Target',      '15.1+  (iPhone only)'),
        ('Version',         '1.0  (Build 1)'),
        ('Date',            str(date.today())),
        ('Status',          '🔴  NOT READY — 8 Critical Blockers'),
        ('Overall Score',   '41 / 100'),
    ]
    for k, v in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, f'{k}:  ', bold=True, size=11, color=C_DARK_GREY)
        add_run(p, v, bold=False, size=11, color=C_MID_GREY)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━', size=12, color=C_GOLD)
    doc.add_page_break()


def section_executive_summary(doc):
    heading(doc, '1  Executive Summary', level=1, color=C_DARK_GREY)
    body(doc,
         'Shazlo is a React Native fashion discovery application currently live on Android (Google Play Store). '
         'This audit covers all code in app/, backend/, and backend/backend/ plus the live production API at '
         'https://api.shazlo.store/ in preparation for iOS App Store submission. '
         'The audit identified 8 critical blockers that must be resolved before any TestFlight or App Store build, '
         '12 high-severity issues, 9 medium issues, and active production data corruption bugs affecting live users.',
         size=10)
    doc.add_paragraph()

    heading(doc, '1.1  Readiness Scores', level=2)
    make_table(doc,
        ['Category', 'Score', 'Status'],
        [
            ['App Architecture',      '55 / 100', '⚠️  Needs Work'],
            ['iOS Native Setup',      '40 / 100', '🔴  Critical Issues'],
            ['Backend Security',      '20 / 100', '🔴  Dangerous'],
            ['Frontend Code Quality', '50 / 100', '⚠️  Needs Work'],
            ['App Store Compliance',  '35 / 100', '🔴  Rejection Risk'],
            ['Privacy & Legal',       '45 / 100', '⚠️  Incomplete'],
            ['Overall',               '41 / 100', '🔴  NOT SUBMITTABLE'],
        ],
        col_widths=[2.5, 1.5, 2.0],
    )

    heading(doc, '1.2  Critical Blocker Summary', level=2)
    make_table(doc,
        ['ID', 'Blocker', 'File', 'Apple Outcome'],
        [
            ['CB-01', 'API_BASE_URL → local LAN IP (HTTP)',       'config/api.js:2',          'App non-functional'],
            ['CB-02', 'GoogleService-Info.plist missing',          'app/ios/',                 'Crash at launch'],
            ['CB-03', 'FirebaseApp.configure() not called',        'AppDelegate.swift',        'Crash at launch'],
            ['CB-04', 'Empty NSLocationWhenInUseUsageDescription', 'Info.plist',               'Automated binary rejection'],
            ['CB-05', 'Sign in with Apple not implemented',        'Onboard.jsx',              'Guideline 4.8 rejection'],
            ['CB-06', 'useSelector inside useEffect',              'Home.tsx:162',             'iOS crash (Rules of Hooks)'],
            ['CB-07', 'DEBUG=True, CORS_ALLOW_ALL_ORIGINS=True',  'settings.py',              'Security / data breach'],
            ['CB-08', 'Zero backend authentication',               'All API views',            'Full API exposure'],
        ],
        col_widths=[0.7, 2.8, 2.0, 1.8],
    )
    doc.add_page_break()


def section_bug_registry(doc):
    heading(doc, '2  Complete Bug Registry', level=1, color=C_DARK_GREY)
    body(doc,
         'Every identified bug is listed below with severity, root cause, and actionable fix steps. '
         'Bugs are ordered by severity. Apply fixes in numerical order.',
         size=10)

    bugs = [
        {
            'id': 'BUG-001',
            'severity': 'CRITICAL',
            'title': 'API_BASE_URL points to local development IP',
            'file': 'app/src/config/api.js',
            'line': '2',
            'symptom': 'Every single network call (login, feed, swipe, cart, orders) fails silently. App appears broken on all non-developer devices.',
            'root_cause': 'Hardcoded local LAN address http://192.168.31.12:8000 committed to source. Apple ATS also blocks HTTP by default.',
            'impact': 'App is entirely non-functional outside developer\'s home network.',
            'fix_steps': [
                '1. Install react-native-config:  npm install react-native-config',
                '2. Create app/.env with:  API_BASE_URL=https://api.shazlo.store',
                '3. Create app/.env.development with:  API_BASE_URL=http://192.168.31.12:8000',
                '4. Add to app/android/app/build.gradle:  apply from: project(\':react-native-config\').projectDir.getPath() + "/dotenv.gradle"',
                '5. Run pod install inside app/ios/',
                '6. Update config/api.js (see code below)',
            ],
            'fix_code': (
                "// app/src/config/api.js  — AFTER FIX\n"
                "import Config from 'react-native-config';\n\n"
                "export const API_BASE_URL = Config.API_BASE_URL;\n\n"
                "export const getImageUrl = (url) =>\n"
                "  `${API_BASE_URL}/v1/items/getimage?url=${encodeURIComponent(url)}`;"
            ),
        },
        {
            'id': 'BUG-002',
            'severity': 'CRITICAL',
            'title': 'GoogleService-Info.plist missing — Firebase cannot initialise on iOS',
            'file': 'app/ios/  (missing file)',
            'line': 'N/A',
            'symptom': 'App crashes at launch on any real iOS device. messaging().getToken() throws immediately.',
            'root_cause': 'The Android equivalent (google-services.json) exists but the iOS config was never downloaded from Firebase Console.',
            'impact': 'Complete crash before first frame renders. App cannot pass TestFlight upload.',
            'fix_steps': [
                '1. Open Firebase Console → Project shazlo-c3712',
                '2. Click Project Settings (gear icon) → Your Apps',
                '3. Add iOS app if not present: Bundle ID = com.shazlo.mobile',
                '4. Download GoogleService-Info.plist',
                '5. In Xcode: drag GoogleService-Info.plist into the Shazlo target (NOT just the folder)',
                '6. Ensure "Copy items if needed" is checked',
                '7. Ensure the file appears in Build Phases → Copy Bundle Resources',
            ],
            'fix_code': (
                "# Terminal — verify file is in Xcode bundle:\n"
                "ls app/ios/Shazlo/GoogleService-Info.plist\n\n"
                "# Expected output:\n"
                "app/ios/Shazlo/GoogleService-Info.plist"
            ),
        },
        {
            'id': 'BUG-003',
            'severity': 'CRITICAL',
            'title': 'FirebaseApp.configure() never called in AppDelegate',
            'file': 'app/ios/Shazlo/AppDelegate.swift',
            'line': 'didFinishLaunchingWithOptions',
            'symptom': 'Even after adding GoogleService-Info.plist, all Firebase SDK calls (messaging, analytics) will throw "Default app has not been configured".',
            'root_cause': 'AppDelegate.swift was not updated after Firebase pods were added.',
            'impact': 'Crash at launch. Push notifications completely broken.',
            'fix_steps': [
                '1. Open app/ios/Shazlo/AppDelegate.swift',
                '2. Add:  import Firebase  at the top',
                '3. Add:  FirebaseApp.configure()  as the very first line of application(_:didFinishLaunchingWithOptions:)',
                '4. Add UNUserNotificationCenterDelegate and MessagingDelegate conformance (see full code below)',
                '5. Clean build folder in Xcode: Product → Clean Build Folder',
                '6. Archive and test on device',
            ],
            'fix_code': (
                "// app/ios/Shazlo/AppDelegate.swift  — FULL REPLACEMENT\n"
                "import UIKit\n"
                "import React\n"
                "import React_RCTAppDelegate\n"
                "import ReactAppDependencyProvider\n"
                "import Firebase\n"
                "import FirebaseMessaging\n"
                "import UserNotifications\n\n"
                "@main\n"
                "class AppDelegate: RCTAppDelegate, UNUserNotificationCenterDelegate, MessagingDelegate {\n\n"
                "  override func application(\n"
                "    _ application: UIApplication,\n"
                "    didFinishLaunchingWithOptions launchOptions: [UIApplication.LaunchOptionsKey: Any]? = nil\n"
                "  ) -> Bool {\n"
                "    FirebaseApp.configure()                    // MUST be first\n"
                "    Messaging.messaging().delegate = self\n"
                "    UNUserNotificationCenter.current().delegate = self\n\n"
                "    self.moduleName = \"app\"\n"
                "    self.dependencyProvider = RCTAppDependencyProvider()\n"
                "    self.initialProps = [:]\n"
                "    return super.application(application, didFinishLaunchingWithOptions: launchOptions)\n"
                "  }\n\n"
                "  func messaging(_ messaging: Messaging,\n"
                "                  didReceiveRegistrationToken fcmToken: String?) {\n"
                "    NotificationCenter.default.post(\n"
                "      name: Notification.Name(\"FCMToken\"),\n"
                "      object: nil,\n"
                "      userInfo: [\"token\": fcmToken ?? \"\"]\n"
                "    )\n"
                "  }\n\n"
                "  override func application(\n"
                "    _ application: UIApplication,\n"
                "    didRegisterForRemoteNotificationsWithDeviceToken deviceToken: Data\n"
                "  ) {\n"
                "    Messaging.messaging().apnsToken = deviceToken\n"
                "  }\n"
                "}"
            ),
        },
        {
            'id': 'BUG-004',
            'severity': 'CRITICAL',
            'title': 'Empty NSLocationWhenInUseUsageDescription — automated binary rejection',
            'file': 'app/ios/Shazlo/Info.plist',
            'line': 'NSLocationWhenInUseUsageDescription',
            'symptom': 'Apple\'s automated review system rejects the binary without human review.',
            'root_cause': 'Empty string "" provided instead of a meaningful user-facing explanation.',
            'impact': 'Guaranteed automatic App Store rejection. Binary never reaches human review.',
            'fix_steps': [
                '1. Open app/ios/Shazlo/Info.plist in Xcode or a text editor',
                '2. Locate the key NSLocationWhenInUseUsageDescription',
                '3. Change its value from "" to a meaningful string',
                '4. Also remove NSAllowsLocalNetworking if present (not needed for production)',
                '5. Rebuild the app',
            ],
            'fix_code': (
                "<!-- app/ios/Shazlo/Info.plist  — CHANGE -->\n\n"
                "<!-- BEFORE (causes rejection): -->\n"
                "<key>NSLocationWhenInUseUsageDescription</key>\n"
                "<string></string>\n\n"
                "<!-- AFTER: -->\n"
                "<key>NSLocationWhenInUseUsageDescription</key>\n"
                "<string>Shazlo uses your location to show you fashion "
                "stores and deals near you.</string>"
            ),
        },
        {
            'id': 'BUG-005',
            'severity': 'CRITICAL',
            'title': 'Sign in with Apple not implemented (Apple Guideline 4.8)',
            'file': 'app/src/screens/Onboard.jsx',
            'line': '156–172',
            'symptom': 'App Store review will reject the app under Guideline 4.8: apps that offer third-party login (Google) MUST also offer Sign in with Apple.',
            'root_cause': 'Google Sign-In SDK is installed and referenced but Apple\'s required counterpart was never implemented.',
            'impact': 'App Store rejection. This is one of the most common rejection reasons.',
            'fix_steps': [
                '1. Run:  npm install @invertase/react-native-apple-authentication',
                '2. In Xcode → Signing & Capabilities → + Capability → Sign In with Apple',
                '3. Create a Services ID in Apple Developer Portal → Identifiers',
                '4. Implement AppleSignInButton component (see code)',
                '5. Call backend /v1/auth/apple-login endpoint (needs backend implementation)',
                '6. Test on real device — Apple auth requires a physical device',
            ],
            'fix_code': (
                "// app/src/components/AppleSignInButton.jsx\n"
                "import React from 'react';\n"
                "import { appleAuth } from '@invertase/react-native-apple-authentication';\n\n"
                "export const signInWithApple = async () => {\n"
                "  const appleAuthResponse = await appleAuth.performRequest({\n"
                "    requestedOperation: appleAuth.Operation.LOGIN,\n"
                "    requestedScopes: [appleAuth.Scope.FULL_NAME, appleAuth.Scope.EMAIL],\n"
                "  });\n"
                "  const { identityToken, fullName } = appleAuthResponse;\n"
                "  // Send identityToken to your backend /v1/auth/apple-login\n"
                "  return { identityToken, fullName };\n"
                "};"
            ),
        },
        {
            'id': 'BUG-006',
            'severity': 'CRITICAL',
            'title': 'useSelector called inside useEffect — React Rules of Hooks violation',
            'file': 'app/src/screens/Home.tsx',
            'line': '162',
            'symptom': 'React throws "Invalid hook call" and crashes the app. The error only manifests on iOS production builds (not Android debug).',
            'root_cause': 'useSelector is a React hook and must be called unconditionally at the top level of a component, not inside useEffect.',
            'impact': 'Hard crash on iOS. Home screen (the primary screen) is unusable.',
            'fix_steps': [
                '1. Open app/src/screens/Home.tsx',
                '2. Find the useSelector call on line 162 inside a useEffect',
                '3. Move it to the top of the component function body (outside all callbacks)',
                '4. Reference the returned value inside the useEffect instead',
            ],
            'fix_code': (
                "// Home.tsx  — BEFORE (crashes):\n"
                "useEffect(() => {\n"
                "  const user = useSelector((state: any) => state.auth.user);  // ❌ INVALID\n"
                "  fetchClosets(user.user_id);\n"
                "}, []);\n\n"
                "// Home.tsx  — AFTER (correct):\n"
                "const user = useSelector((state: any) => state.auth.user);  // ✅ top level\n\n"
                "useEffect(() => {\n"
                "  if (user?.user_id) fetchClosets(user.user_id);\n"
                "}, [user?.user_id]);"
            ),
        },
        {
            'id': 'BUG-007',
            'severity': 'CRITICAL',
            'title': 'Backend: DEBUG=True and CORS_ALLOW_ALL_ORIGINS=True in production',
            'file': 'backend/backend/backend/settings.py',
            'line': '~20–40',
            'symptom': 'Full Django debug page (stack trace + source code + environment variables + DB schema) exposed on any 500 error. Any website can make authenticated API calls.',
            'root_cause': 'Development settings committed and never overridden for production.',
            'impact': 'Complete security failure. Exposes database structure, secret keys, and all internal state.',
            'fix_steps': [
                '1. Set DEBUG=False in Render environment variables',
                '2. Set ALLOWED_HOSTS=api.shazlo.store in Render environment variables',
                '3. Set CORS_ALLOWED_ORIGINS=https://www.shazlo.store,https://shazlo.store in Render env',
                '4. Change CORS_ALLOW_ALL_ORIGINS to False in settings.py',
                '5. Add a custom 500.html error template',
                '6. Deploy and verify error pages no longer show debug info',
            ],
            'fix_code': (
                "# backend/backend/backend/settings.py  — PRODUCTION CHANGES\n\n"
                "DEBUG = config('DEBUG', default=False, cast=bool)  # Render: DEBUG=False\n\n"
                "ALLOWED_HOSTS = config(\n"
                "    'ALLOWED_HOSTS',\n"
                "    default='api.shazlo.store'\n"
                ").split(',')\n\n"
                "CORS_ALLOW_ALL_ORIGINS = False\n"
                "CORS_ALLOWED_ORIGINS = config(\n"
                "    'CORS_ALLOWED_ORIGINS',\n"
                "    default='https://www.shazlo.store'\n"
                ").split(',')"
            ),
        },
        {
            'id': 'BUG-008',
            'severity': 'CRITICAL',
            'title': 'Zero authentication on all API endpoints',
            'file': 'backend/backend/api/views/  (all files)',
            'line': 'All @api_view decorators',
            'symptom': 'Any person who knows any user\'s UUID can read their profile, see their cart, place orders, delete their account, and perform swipe actions on their behalf.',
            'root_cause': 'JWT or session authentication was never implemented. All views accept a user_id parameter and trust it unconditionally.',
            'impact': 'Complete account takeover vulnerability. This is a showstopper security issue.',
            'fix_steps': [
                '1. Add djangorestframework-simplejwt to requirements.txt',
                '2. Configure JWT in settings.py (see code)',
                '3. Issue JWT access+refresh tokens at /v1/auth/shadow and /v1/auth/google-login',
                '4. Add @permission_classes([IsAuthenticated]) to all sensitive views',
                '5. Extract user_id from request.user instead of request.data',
                '6. Implement token refresh endpoint',
                '7. Update React Native app to store and send Authorization header',
            ],
            'fix_code': (
                "# Step 1 — requirements.txt\n"
                "djangorestframework-simplejwt==5.5.0\n\n"
                "# Step 2 — settings.py\n"
                "REST_FRAMEWORK = {\n"
                "    'DEFAULT_AUTHENTICATION_CLASSES': (\n"
                "        'rest_framework_simplejwt.authentication.JWTAuthentication',\n"
                "    ),\n"
                "}\n\n"
                "# Step 3 — shadow user login response\n"
                "from rest_framework_simplejwt.tokens import RefreshToken\n\n"
                "refresh = RefreshToken.for_user(user)\n"
                "return Response({\n"
                "    'user': {...},\n"
                "    'access':  str(refresh.access_token),\n"
                "    'refresh': str(refresh),\n"
                "}, status=201)\n\n"
                "# Step 4 — protected view example\n"
                "@api_view(['GET'])\n"
                "@permission_classes([IsAuthenticated])\n"
                "def get_cart(request):\n"
                "    user_id = request.user.user_id   # from JWT, not request.data"
            ),
        },
        {
            'id': 'BUG-009',
            'severity': 'HIGH',
            'title': 'Active data corruption: user.email overwritten with phone_number',
            'file': 'backend/backend/api/views/user_views.py',
            'line': '245',
            'symptom': 'After any user updates their profile, their email field contains their phone number. Login via email breaks for affected users.',
            'root_cause': 'Typo: user.email = phone_number instead of user.email = email.',
            'impact': 'Production data is being actively corrupted for every user who completes signup.',
            'fix_steps': [
                '1. Open backend/backend/api/views/user_views.py',
                '2. Find line 245: user.email = phone_number',
                '3. Change to: user.email = email',
                '4. Deploy immediately',
                '5. Run a one-time data repair script to restore emails from Action logs or re-registration',
            ],
            'fix_code': (
                "# user_views.py line 245  — CHANGE ONE WORD\n\n"
                "# BEFORE (data corruption):\n"
                "user.email = phone_number   # ❌ BUG\n\n"
                "# AFTER:\n"
                "user.email = email          # ✅ correct"
            ),
        },
        {
            'id': 'BUG-010',
            'severity': 'HIGH',
            'title': 'Raw passwords and PII printed to production logs',
            'file': 'backend/backend/api/views/user_views.py',
            'line': '116, 219, 302',
            'symptom': 'Render dashboard logs expose plaintext passwords, phone numbers, and emails for every login and registration event.',
            'root_cause': 'Debug print statements left in production code.',
            'impact': 'Credential leakage. Violates GDPR Article 32 and App Store privacy requirements.',
            'fix_steps': [
                '1. Remove line 302: print(identifier, password)',
                '2. Remove line 219: print(phone_number)',
                '3. Remove line 116: print(phone_number)',
                '4. Search entire codebase: grep -rn "print(" backend/ and remove all debug prints',
                '5. Add structured logging with log level controls instead',
            ],
            'fix_code': (
                "# user_views.py — REMOVE these lines immediately:\n"
                "# Line 302:  print(identifier, password)   ← REMOVE\n"
                "# Line 219:  print(phone_number)            ← REMOVE\n"
                "# Line 116:  print(phone_number)            ← REMOVE\n\n"
                "# If logging is needed, use Python's logging module:\n"
                "import logging\n"
                "logger = logging.getLogger(__name__)\n"
                "logger.info('Login attempt for identifier: %s', identifier[:3] + '***')"
            ),
        },
        {
            'id': 'BUG-011',
            'severity': 'HIGH',
            'title': 'SSRF vulnerability in image proxy endpoint',
            'file': 'backend/backend/api/views/item_views.py',
            'line': 'proxy_image (GET /v1/items/getimage)',
            'symptom': 'An attacker can call /v1/items/getimage?url=http://169.254.169.254/latest/meta-data/ to read AWS/cloud IMDS and steal IAM credentials.',
            'root_cause': 'Open proxy with no URL validation, domain allowlist, or request timeout.',
            'impact': 'Cloud infrastructure compromise. Attacker can steal API keys, DB credentials, and service tokens.',
            'fix_steps': [
                '1. Add a domain allowlist (only allow Cloudinary CDN)',
                '2. Block private/reserved IP ranges before making request',
                '3. Set a short request timeout (5 seconds)',
                '4. Migrate all images to Cloudinary direct URLs (eliminates need for proxy entirely)',
            ],
            'fix_code': (
                "# item_views.py — SSRF-safe proxy\n"
                "import ipaddress, socket\n"
                "from urllib.parse import urlparse\n\n"
                "ALLOWED_DOMAINS = {'res.cloudinary.com', 'images.shazlo.store'}\n\n"
                "def proxy_image(request):\n"
                "    url = request.GET.get('url', '')\n"
                "    parsed = urlparse(url)\n"
                "    if parsed.hostname not in ALLOWED_DOMAINS:\n"
                "        return HttpResponse(status=403)\n"
                "    # Additional: resolve hostname and block private IPs\n"
                "    try:\n"
                "        ip = socket.gethostbyname(parsed.hostname)\n"
                "        if ipaddress.ip_address(ip).is_private:\n"
                "            return HttpResponse(status=403)\n"
                "    except Exception:\n"
                "        return HttpResponse(status=403)\n"
                "    resp = requests.get(url, timeout=5, stream=True)\n"
                "    return StreamingHttpResponse(resp.iter_content(8192),\n"
                "                               content_type=resp.headers.get('Content-Type'))"
            ),
        },
        {
            'id': 'BUG-012',
            'severity': 'HIGH',
            'title': 'Duplicate NetInfo event listeners — memory leak',
            'file': 'app/App.jsx',
            'line': '40–58',
            'symptom': 'Memory usage grows continuously. Two offline detection banners may appear simultaneously.',
            'root_cause': 'Two separate useEffect blocks both call NetInfo.addEventListener without cleanup.',
            'impact': 'Memory leak. Incorrect offline state. Performance degradation.',
            'fix_steps': [
                '1. Open app/App.jsx',
                '2. Find the two useEffect blocks both calling NetInfo.addEventListener',
                '3. Merge them into a single useEffect',
                '4. Return the unsubscribe function from the merged useEffect',
            ],
            'fix_code': (
                "// App.jsx — merge the two NetInfo useEffects:\n"
                "useEffect(() => {\n"
                "  const unsubscribe = NetInfo.addEventListener(state => {\n"
                "    setIsOffline(!state.isConnected);\n"
                "    if (state.isConnected) hideOfflineBanner();\n"
                "    else showOfflineBanner();\n"
                "  });\n"
                "  return () => unsubscribe();   // cleanup on unmount\n"
                "}, []);"
            ),
        },
        {
            'id': 'BUG-013',
            'severity': 'HIGH',
            'title': 'QueryClient created inside App component — destroyed on every re-render',
            'file': 'app/App.jsx',
            'line': '136',
            'symptom': 'All React Query cached data (cart, recommendations) is wiped every time App re-renders. Users see loading spinners unexpectedly.',
            'root_cause': 'new QueryClient() called inside the component body, creating a new instance on each render cycle.',
            'impact': 'Cart data loss, poor UX, excess network requests.',
            'fix_steps': [
                '1. Open app/App.jsx',
                '2. Find: const queryClient = new QueryClient() inside the App function',
                '3. Move that line to the module level — outside the App function entirely',
            ],
            'fix_code': (
                "// App.jsx — BEFORE (wrong):\n"
                "const App = () => {\n"
                "  const queryClient = new QueryClient();   // ❌ recreated every render\n"
                "  return <QueryClientProvider client={queryClient}>...</QueryClientProvider>;\n"
                "};\n\n"
                "// App.jsx — AFTER (correct):\n"
                "const queryClient = new QueryClient();    // ✅ module-level singleton\n\n"
                "const App = () => {\n"
                "  return <QueryClientProvider client={queryClient}>...</QueryClientProvider>;\n"
                "};"
            ),
        },
        {
            'id': 'BUG-014',
            'severity': 'HIGH',
            'title': 'PII sent in URL query parameters in Payment screen',
            'file': 'app/src/screens/Payment.jsx',
            'line': '73',
            'symptom': 'User name and phone number appear in server access logs, browser history, and referrer headers.',
            'root_cause': 'Payment URL constructed with name=${user.name}&mno=${user.phone_number} as query params.',
            'impact': 'PII leakage. Violates GDPR, Apple App Store privacy policy requirements.',
            'fix_steps': [
                '1. Open app/src/screens/Payment.jsx',
                '2. Change the checkout URL to use POST body or load the Razorpay SDK natively',
                '3. Send user data as POST body to backend, which constructs the Razorpay order server-side',
                '4. Never include PII in URLs',
            ],
            'fix_code': (
                "// Payment.jsx — send PII via POST body, not URL:\n"
                "const createOrder = async () => {\n"
                "  const response = await fetch(`${API_BASE_URL}/v1/payments/create-order`, {\n"
                "    method: 'POST',\n"
                "    headers: { 'Content-Type': 'application/json' },\n"
                "    body: JSON.stringify({\n"
                "      amount: totalAmount,\n"
                "      user_id: user.user_id,  // eventually replaced by JWT\n"
                "    }),\n"
                "  });\n"
                "  const order = await response.json();\n"
                "  // Open Razorpay with order.id\n"
                "};"
            ),
        },
        {
            'id': 'BUG-015',
            'severity': 'MEDIUM',
            'title': 'Hardcoded placeholder product details on back-of-card for all items',
            'file': 'app/src/screens/SwipeUI.jsx',
            'line': '1688–1706',
            'symptom': 'Every single product card shows "100% premium cotton", "Machine wash cold", "Easy 7-day return policy" regardless of the actual product.',
            'root_cause': 'Static placeholder copy was never replaced with real product data from the API.',
            'impact': 'False product information displayed. Potential Apple guideline 2.3 (Accurate Metadata). Consumer trust issue.',
            'fix_steps': [
                '1. Add material, care_instructions, and return_policy fields to the Item model',
                '2. Populate those fields during scraping/ingestion',
                '3. Remove hardcoded strings in SwipeUI.jsx lines 1688–1706',
                '4. Display item.material, item.care_instructions from the API response',
                '5. If fields are empty, hide the section entirely rather than showing placeholder',
            ],
            'fix_code': (
                "// SwipeUI.jsx — replace hardcoded placeholders:\n"
                "// BEFORE:\n"
                "<Text>100% premium cotton</Text>\n"
                "<Text>Machine wash cold</Text>\n\n"
                "// AFTER:\n"
                "{item.material && <Text>{item.material}</Text>}\n"
                "{item.care_instructions && <Text>{item.care_instructions}</Text>}"
            ),
        },
        {
            'id': 'BUG-016',
            'severity': 'MEDIUM',
            'title': 'Admin endpoint exposed publicly with no authentication',
            'file': 'backend/backend/api/views/recommendation_views.py',
            'line': 'find_duplicate_images',
            'symptom': 'Anyone can call /v1/recommendations/find-duplicates to trigger a database scan that deletes items and modifies all user carts.',
            'root_cause': 'Administrative endpoint was built but never restricted.',
            'impact': 'Data destruction. Cart data corruption. DoS via repeated calls.',
            'fix_steps': [
                '1. Add @permission_classes([IsAdminUser]) decorator to find_duplicate_images view',
                '2. Or remove the endpoint and run it as a management command: python manage.py deduplicate_images',
                '3. Ensure staff users are properly configured in Django admin',
            ],
            'fix_code': (
                "# recommendation_views.py\n"
                "from rest_framework.permissions import IsAdminUser\n\n"
                "@api_view(['POST'])\n"
                "@permission_classes([IsAdminUser])   # ← add this\n"
                "def find_duplicate_images(request):\n"
                "    ..."
            ),
        },
        {
            'id': 'BUG-017',
            'severity': 'MEDIUM',
            'title': 'Unencrypted storage of full user profile in AsyncStorage',
            'file': 'app/src/store/index.ts',
            'line': 'redux-persist config',
            'symptom': 'User name, email, phone number, gender, date of birth, and 512-float preference vector stored in plaintext on device.',
            'root_cause': 'redux-persist uses AsyncStorage by default with no encryption.',
            'impact': 'PII accessible to other apps on jailbroken devices. Fails Apple data security requirements.',
            'fix_steps': [
                '1. Replace redux-persist\'s AsyncStorage engine with MMKV encrypted storage',
                '2. Generate a secure encryption key (use react-native-keychain to store it)',
                '3. Or exclude sensitive fields (email, phone, DOB) from persistence',
                '4. Only persist: user_id, name, is_shadow, preference_vector (not PII)',
            ],
            'fix_code': (
                "// store/index.ts — use encrypted MMKV storage\n"
                "import { MMKV } from 'react-native-mmkv';\n\n"
                "const storage = new MMKV({ id: 'redux-store', encryptionKey: 'your-secure-key' });\n\n"
                "const mmkvStorage = {\n"
                "  setItem: (key, value) => storage.set(key, value),\n"
                "  getItem: (key) => storage.getString(key) ?? null,\n"
                "  removeItem: (key) => storage.delete(key),\n"
                "};\n\n"
                "const persistConfig = {\n"
                "  key: 'root',\n"
                "  storage: mmkvStorage,\n"
                "  whitelist: ['auth'],  // only persist auth\n"
                "};"
            ),
        },
        {
            'id': 'BUG-018',
            'severity': 'MEDIUM',
            'title': 'console.log statements log entire item arrays on every render',
            'file': 'app/src/screens/SwipeUI.jsx',
            'line': '73, 1054',
            'symptom': 'Thousands of items logged to Metro bundler and Xcode console on every swipe. Performance impact in production.',
            'root_cause': 'Debug logs committed and never stripped.',
            'impact': 'Performance degradation. Log spam. Sensitive product data in device logs.',
            'fix_steps': [
                '1. Add babel-plugin-transform-remove-console to strip all console.* in production',
                '2. Or manually remove console.log(items) and console.log(brand)',
                '3. Configure in babel.config.js (see code)',
            ],
            'fix_code': (
                "// babel.config.js — strip console.log in production builds\n"
                "module.exports = {\n"
                "  presets: ['module:@react-native/babel-preset'],\n"
                "  env: {\n"
                "    production: {\n"
                "      plugins: ['transform-remove-console'],\n"
                "    },\n"
                "  },\n"
                "};"
            ),
        },
        {
            'id': 'BUG-019',
            'severity': 'MEDIUM',
            'title': 'Google Sign-In not wired to backend — orphaned implementation',
            'file': 'app/src/screens/Onboard.jsx',
            'line': '156–172',
            'symptom': 'Google Sign-In button is commented out. Even when uncommented, onGoogleButtonPress never calls the backend or dispatches setlogin.',
            'root_cause': 'Partial implementation — only the Google auth SDK call was written, not the backend integration or Redux dispatch.',
            'impact': 'Feature non-functional. Required for Sign in with Apple parity (Guideline 4.8).',
            'fix_steps': [
                '1. Uncomment the GoogleButton in Onboard.jsx',
                '2. In onGoogleButtonPress, after obtaining idToken, POST to /v1/auth/google-login',
                '3. Dispatch setlogin with the returned user object',
                '4. Handle navigation to Home or Tutorial after successful login',
                '5. Move the hardcoded webClientId to react-native-config (.env)',
            ],
            'fix_code': (
                "// Onboard.jsx — complete Google Sign-In flow\n"
                "const onGoogleButtonPress = async () => {\n"
                "  await GoogleSignin.hasPlayServices();\n"
                "  const userInfo = await GoogleSignin.signIn();\n"
                "  const idToken = userInfo.data?.idToken;\n\n"
                "  const res = await fetch(`${API_BASE_URL}/v1/auth/google-login`, {\n"
                "    method: 'POST',\n"
                "    headers: { 'Content-Type': 'application/json' },\n"
                "    body: JSON.stringify({ id_token: idToken }),\n"
                "  });\n"
                "  const data = await res.json();\n"
                "  if (res.ok) {\n"
                "    dispatch(setlogin({ user: data.user }));\n"
                "    navigation.dispatch(CommonActions.reset({ index: 0, routes: [{ name: 'Home' }] }));\n"
                "  }\n"
                "};"
            ),
        },
        {
            'id': 'BUG-020',
            'severity': 'MEDIUM',
            'title': 'Landscape orientation declared but UI is portrait-only',
            'file': 'app/ios/Shazlo/Info.plist',
            'line': 'UISupportedInterfaceOrientations',
            'symptom': 'App can rotate to landscape on iPad and iPhone, breaking all UI which was designed for portrait.',
            'root_cause': 'Default Info.plist includes LandscapeLeft and LandscapeRight which was never removed.',
            'impact': 'UI breaks in landscape. Poor UX during App Store review.',
            'fix_steps': [
                '1. Open Info.plist',
                '2. Under UISupportedInterfaceOrientations, remove UIInterfaceOrientationLandscapeLeft',
                '3. Remove UIInterfaceOrientationLandscapeRight',
                '4. Keep only UIInterfaceOrientationPortrait',
            ],
            'fix_code': (
                "<!-- Info.plist — portrait only -->\n"
                "<key>UISupportedInterfaceOrientations</key>\n"
                "<array>\n"
                "  <string>UIInterfaceOrientationPortrait</string>\n"
                "  <!-- Remove landscape entries below -->\n"
                "  <!-- <string>UIInterfaceOrientationLandscapeLeft</string>  ← DELETE -->\n"
                "  <!-- <string>UIInterfaceOrientationLandscapeRight</string> ← DELETE -->\n"
                "</array>"
            ),
        },
    ]

    for bug in bugs:
        sev = bug['severity']
        sev_colors = {
            'CRITICAL': (C_RED, 'D32F2F'),
            'HIGH':     (C_ORANGE, 'F57C00'),
            'MEDIUM':   (C_YELLOW, 'F9A825'),
            'LOW':      (C_GREEN, '2E7D32'),
        }
        text_color, bg_hex = sev_colors.get(sev, (C_MID_GREY, 'AAAAAA'))

        heading(doc, f"{bug['id']}  —  {bug['title']}", level=2, color=C_DARK_GREY)

        # Metadata bar
        p = doc.add_paragraph()
        add_run(p, f"  {severity_badge(sev)}  ", bold=True, size=9.5, color=text_color)
        add_run(p, f"  |  File: ", bold=True, size=9, color=C_MID_GREY)
        add_run(p, f"{bug['file']}  line {bug['line']}", bold=False, size=9, color=C_BLUE)

        body(doc, f"Symptom:    {bug['symptom']}", size=9.5)
        body(doc, f"Root Cause: {bug['root_cause']}", size=9.5)
        body(doc, f"Impact:     {bug['impact']}", size=9.5)

        heading(doc, 'Fix Steps:', level=4, color=C_DARK_GREY)
        for step in bug['fix_steps']:
            bullet(doc, step, size=9)

        heading(doc, 'Code:', level=4, color=C_DARK_GREY)
        for line in bug['fix_code'].split('\n'):
            code_block(doc, line)

        divider(doc)

    doc.add_page_break()


def section_ios_deployment(doc):
    heading(doc, '3  iOS Setup & Deployment — Step-by-Step', level=1, color=C_DARK_GREY)
    body(doc, 'Follow every step in order. Steps are grouped into phases. Do not skip any step.')

    phases = [
        {
            'title': 'Phase 1  —  Prerequisites & Environment (Day 1)',
            'steps': [
                ('macOS machine', 'Required. Xcode cannot run on any other OS.'),
                ('Xcode 16.x', 'Download from Mac App Store. Select "iOS 18 SDK" during install.'),
                ('Apple Developer Account', 'Enroll at developer.apple.com. Individual: $99/year. Takes up to 48 hours to activate.'),
                ('Node 18+ and CocoaPods', 'brew install node; sudo gem install cocoapods'),
                ('Ruby 3.x',  'rbenv install 3.2.0 && rbenv global 3.2.0  (avoids system Ruby conflicts)'),
                ('Watchman',  'brew install watchman'),
                ('React Native CLI', 'npm install -g @react-native-community/cli'),
            ]
        },
        {
            'title': 'Phase 2  —  Critical Fixes (Day 1–2)',
            'steps': [
                ('Fix API_BASE_URL', 'Apply BUG-001 fix. All network calls depend on this.'),
                ('Fix Home.tsx hook violation', 'Apply BUG-006 fix before any iOS build attempt.'),
                ('Fix user_views.py data corruption', 'Apply BUG-009 fix and deploy backend immediately.'),
                ('Remove debug print statements', 'Apply BUG-010 fix and redeploy backend.'),
                ('Fix Django production settings', 'Apply BUG-007 fix. Set DEBUG=False in Render dashboard.'),
            ]
        },
        {
            'title': 'Phase 3  —  Firebase iOS Setup (Day 2)',
            'steps': [
                ('Firebase Console', 'Go to console.firebase.google.com → Project shazlo-c3712'),
                ('Add iOS app', 'Project Settings → Your Apps → Add App → iOS. Bundle ID: com.shazlo.mobile'),
                ('Download plist', 'Download GoogleService-Info.plist'),
                ('Add to Xcode', 'Drag GoogleService-Info.plist into Xcode under the Shazlo target. Check "Copy items if needed"'),
                ('Verify in Build Phases', 'Build Phases → Copy Bundle Resources → confirm GoogleService-Info.plist is listed'),
                ('Update AppDelegate', 'Apply BUG-003 fix — full AppDelegate.swift replacement'),
            ]
        },
        {
            'title': 'Phase 4  —  iOS Project Setup (Day 2–3)',
            'steps': [
                ('Open workspace', 'Always open Shazlo.xcworkspace (not .xcodeproj) in Xcode'),
                ('Set team', 'Signing & Capabilities → Team → select your Apple Developer team'),
                ('Verify bundle ID', 'com.shazlo.mobile — must match Firebase and App Store Connect'),
                ('Fix Info.plist', 'Apply BUG-004 and BUG-020 fixes (location string + portrait only)'),
                ('Remove local networking', 'Remove NSAllowsLocalNetworking from Info.plist'),
                ('Add Sign in with Apple capability', 'Signing & Capabilities → + → Sign In with Apple'),
                ('Add push notifications capability', 'Signing & Capabilities → + → Push Notifications'),
                ('Add associated domains', 'Signing & Capabilities → + → Associated Domains → applinks:www.shazlo.store'),
                ('Run pod install', 'cd app/ios && pod install --repo-update'),
            ]
        },
        {
            'title': 'Phase 5  —  APNs (Push Notifications) Setup (Day 3)',
            'steps': [
                ('Create APNs key', 'developer.apple.com → Certificates → Keys → + → Apple Push Notifications service (APNs)'),
                ('Download .p8 file', 'Save AuthKey_XXXXXXXXXX.p8  — can only be downloaded once'),
                ('Upload to Firebase', 'Firebase Console → Project Settings → Cloud Messaging → APNs Authentication Key → Upload'),
                ('Enter Key ID and Team ID', 'Both found in Apple Developer Portal → Keys and Membership'),
                ('Test on device', 'Install a debug build on a real iPhone and verify push notification arrives'),
            ]
        },
        {
            'title': 'Phase 6  —  Sign in with Apple (Day 3–4)',
            'steps': [
                ('Install package', 'npm install @invertase/react-native-apple-authentication && cd ios && pod install'),
                ('Add capability', 'Xcode → Signing & Capabilities → Sign In with Apple (already done in Phase 4)'),
                ('Apple Developer Portal', 'Identifiers → your App ID → enable Sign In with Apple'),
                ('Implement component', 'Apply BUG-005 fix code — create AppleSignInButton.jsx'),
                ('Backend endpoint', 'Create /v1/auth/apple-login view that verifies Apple identity token'),
                ('Test', 'Must test on real device — Apple auth requires physical hardware'),
            ]
        },
        {
            'title': 'Phase 7  —  Complete Google Sign-In (Day 4)',
            'steps': [
                ('Apply BUG-019 fix', 'Wire onGoogleButtonPress to backend and Redux'),
                ('Move client ID to .env', 'GOOGLE_WEB_CLIENT_ID=155811039707-... in app/.env'),
                ('Google Developer Console', 'Add iOS bundle ID com.shazlo.mobile to the OAuth 2.0 client'),
                ('Add URL scheme', 'Info.plist → URL Types → add REVERSED_CLIENT_ID from GoogleService-Info.plist'),
                ('Uncomment button', 'Remove the comment wrapping <GoogleButton> in Onboard.jsx'),
                ('Test end-to-end', 'Google Sign-In → backend login → Home screen navigation'),
            ]
        },
        {
            'title': 'Phase 8  —  TestFlight Build (Day 4–5)',
            'steps': [
                ('Set version', 'Xcode → General → Version: 1.0, Build: 1'),
                ('Select device', 'Set scheme to "Shazlo" and destination to "Any iOS Device (arm64)"'),
                ('Archive', 'Product → Archive  (takes 5–10 minutes)'),
                ('Distribute', 'In Organizer → Distribute App → App Store Connect → Upload'),
                ('Wait for processing', 'Apple processes the build — usually 15–30 minutes'),
                ('Add to TestFlight', 'App Store Connect → TestFlight → select build → Add External Testers'),
                ('Submit for Beta Review', 'Required for external testers — usually 1 business day'),
            ]
        },
        {
            'title': 'Phase 9  —  App Store Submission (Day 5–7)',
            'steps': [
                ('Screenshots', 'Required sizes: 6.9" iPhone (1320×2868), 6.5" iPhone (1242×2688). Use Xcode Simulator or real devices.'),
                ('App description', 'Write in App Store Connect → App Information → Description (min 4 sentences)'),
                ('Privacy policy URL', 'Must be publicly accessible. Host at https://www.shazlo.store/privacy'),
                ('Age rating', 'App Store Connect → Age Rating → answer questionnaire (likely 4+)'),
                ('Category', 'Shopping (primary)'),
                ('Keywords', 'fashion, shopping, swipe, discovery, style, clothing, outfits (max 100 chars)'),
                ('Review notes', 'In "Notes for Reviewer" explain shadow user flow — reviewer needs to be able to use the app without creating account'),
                ('Submit for review', 'App Store Connect → Submit for Review. Standard review: 24–48 hours.'),
            ]
        },
    ]

    for phase in phases:
        heading(doc, phase['title'], level=2, color=C_DARK_GREY)
        make_table(doc,
            ['Step', 'Action', 'Detail'],
            [[str(i+1), s[0], s[1]] for i, s in enumerate(phase['steps'])],
            col_widths=[0.4, 1.8, 4.0],
        )

    doc.add_page_break()


def section_backend_hardening(doc):
    heading(doc, '4  Backend Security Hardening', level=1, color=C_DARK_GREY)
    body(doc, 'Ordered by priority. Steps 1–5 must be completed before any iOS submission.', size=10)

    steps = [
        ('1', 'HIGH',     'Fix DEBUG and CORS',            'Set DEBUG=False, ALLOWED_HOSTS=api.shazlo.store, CORS_ALLOWED_ORIGINS in Render env vars. See BUG-007.'),
        ('2', 'CRITICAL', 'Remove credential logging',     'Delete print(identifier, password) and all PII print statements. See BUG-010.'),
        ('3', 'CRITICAL', 'Implement JWT authentication',  'Install djangorestframework-simplejwt. Issue tokens at login endpoints. Protect all views. See BUG-008.'),
        ('4', 'CRITICAL', 'Fix SSRF proxy',                'Add domain allowlist to /v1/items/getimage. Block private IPs. See BUG-011.'),
        ('5', 'HIGH',     'Restrict admin endpoint',       'Add @permission_classes([IsAdminUser]) to find_duplicate_images. See BUG-016.'),
        ('6', 'HIGH',     'Add rate limiting',             'pip install django-ratelimit. Add @ratelimit(key="ip", rate="10/m") to shadow user and login endpoints.'),
        ('7', 'HIGH',     'Enable HTTPS enforcement',      'Add SECURE_SSL_REDIRECT=True, SECURE_HSTS_SECONDS=31536000 to settings.py. Render provides HTTPS automatically.'),
        ('8', 'MEDIUM',   'Add database connection pooling','DB_POOL_MODE is configured — ensure PgBouncer is enabled in Render PostgreSQL settings.'),
        ('9', 'MEDIUM',   'Implement request logging',     'Use django-request-logging or structlog. Log request method, path, user_id, status — never PII.'),
        ('10','MEDIUM',   'Add Sentry error tracking',     'pip install sentry-sdk. Add SENTRY_DSN to Render env vars. Never log raw user data.'),
    ]

    make_table(doc,
        ['#', 'Priority', 'Task', 'Details'],
        steps,
        col_widths=[0.3, 0.9, 1.8, 3.8],
    )

    heading(doc, '4.1  Rate Limiting Setup', level=2)
    code_block(doc, '# requirements.txt')
    code_block(doc, 'django-ratelimit==4.1.0')
    code_block(doc, '')
    code_block(doc, '# views/user_views.py')
    code_block(doc, 'from ratelimit.decorators import ratelimit')
    code_block(doc, '')
    code_block(doc, '@api_view([\'POST\'])')
    code_block(doc, '@ratelimit(key=\'ip\', rate=\'5/m\', block=True)')
    code_block(doc, 'def create_shadow_user(request):')
    code_block(doc, '    ...')
    doc.add_paragraph()

    heading(doc, '4.2  Render Environment Variables to Set', level=2)
    make_table(doc,
        ['Variable', 'Value', 'Notes'],
        [
            ['DEBUG',                'False',                        'Critical — must be False'],
            ['ALLOWED_HOSTS',        'api.shazlo.store',             'No wildcards'],
            ['CORS_ALLOWED_ORIGINS', 'https://www.shazlo.store',     'Comma-separated if multiple'],
            ['SECRET_KEY',           '(generate 50-char random)',    'Use: python -c "import secrets; print(secrets.token_hex(25))"'],
            ['DATABASE_URL',         '(from Render PostgreSQL)',     'Auto-provided by Render'],
            ['SENTRY_DSN',           '(from sentry.io)',             'For error tracking'],
            ['CLOUDINARY_URL',       '(from Cloudinary dashboard)',  'For image CDN'],
        ],
        col_widths=[2.0, 2.2, 2.5],
    )
    doc.add_page_break()


def section_costing(doc):
    heading(doc, '5  Cost & Infrastructure Analysis', level=1, color=C_DARK_GREY)
    body(doc,
         'All costs are in USD. Figures are approximate as of May 2026. '
         'Shazlo is currently on free tiers for most services — these are the upgrade costs as the app scales.',
         size=10)

    heading(doc, '5.1  One-Time Costs', level=2)
    make_table(doc,
        ['Item', 'Cost (USD)', 'Notes'],
        [
            ['Apple Developer Program (Year 1)',   '$99.00',   'Required to submit to App Store. Renews annually.'],
            ['Google Play Console',                '$25.00',   'One-time. Already paid (app is live on Android).'],
            ['App Icon Design (if outsourced)',    '$50–200',  'If using a designer. Free with Figma/Canva.'],
            ['Privacy Policy page (if outsourced)','$0–100',  'Required. Free via iubenda.com or termly.io.'],
            ['App Screenshots (designer)',         '$0–150',   'Can be generated from Xcode Simulator for free.'],
            ['Code signing certificate',           '$0',       'Included in Apple Developer Program.'],
        ],
        col_widths=[2.8, 1.3, 3.2],
    )

    heading(doc, '5.2  Monthly Recurring Costs (Current Free Tier)', level=2)
    make_table(doc,
        ['Service', 'Current Plan', 'Monthly Cost', 'Limitation'],
        [
            ['Render (Backend)',   'Free',         '$0',    '90s cold start, sleeps after 15 min inactivity'],
            ['Render PostgreSQL',  'Free',         '$0',    '1GB storage, 97 day expiry'],
            ['Firebase',          'Spark (Free)', '$0',    '10K/month Crashlytics, 1M/month Analytics events'],
            ['Cloudinary',        'Free',         '$0',    '25 credits/month, 25GB storage, 25GB bandwidth'],
            ['Supabase',          'Free',         '$0',    '500MB DB, 1GB file storage (ML model)'],
            ['Google APIs',       'Free',         '$0',    'Sign-In quota: 10K requests/day'],
            ['Apple Developer',   'Standard',     '$8.25', 'Billed annually at $99/year'],
            ['Total (Current)',   '—',            '$8.25', 'Per month while on free tiers'],
        ],
        col_widths=[1.8, 1.4, 1.2, 3.0],
    )

    heading(doc, '5.3  Recommended Production Upgrade Costs', level=2)
    body(doc, 'Upgrade at the following user milestones to maintain performance and reliability.', size=9.5)
    make_table(doc,
        ['Service', 'Plan', 'Cost/mo', 'When to Upgrade', 'What You Get'],
        [
            ['Render (Web Service)',   'Starter',       '$7',      '1,000 users',        'No cold start, 512MB RAM, 0.5 CPU'],
            ['Render (Web Service)',   'Standard',      '$25',     '5,000 users',        '2GB RAM, 1 CPU, always-on'],
            ['Render PostgreSQL',      'Basic',         '$7',      '1,000 users',        '1GB RAM, 256GB storage, no expiry'],
            ['Render PostgreSQL',      'Standard',      '$20',     '10,000 users',       '2GB RAM, 1TB storage, read replicas'],
            ['Cloudinary',             'Plus',          '$89',     '10K active users',   '225 credits, 225GB storage+bandwidth'],
            ['Firebase',               'Blaze (Pay-go)','~$5–20', '50K+ events/month',  'Unlimited events, pay per use'],
            ['Supabase',               'Pro',           '$25',     'Model file >1GB',    '8GB DB, 100GB storage'],
            ['Sentry',                 'Team',          '$26',     'iOS launch',         '50K errors/month, alerting, trends'],
        ],
        col_widths=[1.7, 1.2, 0.9, 1.5, 2.0],
    )

    heading(doc, '5.4  Cost Projection by Scale', level=2)
    make_table(doc,
        ['Active Users', 'Render', 'Firebase', 'Cloudinary', 'Postgres', 'Other', 'Total/mo'],
        [
            ['0–1,000',    '$0',   '$0',    '$0',   '$0',   '$8.25',  '$8.25'],
            ['1,000–5,000','$7',   '$0',    '$0',   '$7',   '$8.25',  '$22.25'],
            ['5,000–20K',  '$25',  '$5',    '$89',  '$20',  '$34.25', '$173.25'],
            ['20K–100K',   '$85',  '$20',   '$89',  '$50',  '$60',    '$304'],
            ['100K+',      '$200+','$50+',  '$200+','$100+','$100+',  '$650+'],
        ],
        col_widths=[1.3, 0.9, 1.0, 1.1, 1.0, 0.9, 1.0],
    )

    heading(doc, '5.5  Apple Developer Account Annual Renewal', level=2)
    make_table(doc,
        ['Item', 'Annual Cost', 'Notes'],
        [
            ['Apple Developer Program',  '$99 USD',    'Renews every 12 months. App stays on store during lapse but new builds cannot be submitted.'],
            ['Apple Developer Enterprise','$299 USD',   'Only for in-house app distribution — NOT needed for App Store.'],
            ['TestFlight (included)',     '$0',          'Up to 10,000 external beta testers, 90-day build expiry.'],
        ],
        col_widths=[2.2, 1.3, 3.3],
    )

    heading(doc, '5.6  Third-Party SDK Cost Summary', level=2)
    make_table(doc,
        ['SDK', 'Version', 'License', 'Cost', 'Notes'],
        [
            ['React Native',           '0.80.0',  'MIT',       'Free', 'Open source'],
            ['Firebase SDK (RN)',       '23.5.0',  'Apache 2',  'Free', 'Firebase services billed separately'],
            ['Google Sign-In',          '16.1.2',  'Apache 2',  'Free', 'No per-request cost'],
            ['Apple Sign-In',           'OS built-in','Apple',  'Free', 'Required — no cost'],
            ['Razorpay',                'Latest',  'Proprietary','~2% per transaction', 'Indian payment gateway'],
            ['Redux Toolkit',           'Latest',  'MIT',       'Free', ''],
            ['React Query v5',          '5.85.5',  'MIT',       'Free', ''],
            ['MMKV',                    '2.10.2',  'MIT',       'Free', ''],
            ['react-native-config',     'Latest',  'MIT',       'Free', 'Needs to be added'],
            ['invertase Apple Auth',    'Latest',  'Apache 2',  'Free', 'Needs to be added'],
        ],
        col_widths=[2.0, 1.1, 1.1, 1.4, 1.8],
    )
    doc.add_page_break()


def section_timeline(doc):
    heading(doc, '6  Implementation Timeline', level=1, color=C_DARK_GREY)
    body(doc, 'Estimated timeline for a single engineer working full-time. Parallel tasks are grouped by day.', size=10)

    make_table(doc,
        ['Day', 'Phase', 'Tasks', 'Est. Hours'],
        [
            ['1',   'Critical Fixes',        'BUG-001 (API URL), BUG-006 (Hooks), BUG-009 (email bug), BUG-010 (logs), BUG-007 (Django settings)', '6–8h'],
            ['2',   'Firebase iOS',          'Download GoogleService-Info.plist, add to Xcode, rewrite AppDelegate.swift, pod install, verify launch', '3–4h'],
            ['2',   'Info.plist Fixes',      'BUG-004 (location string), BUG-020 (portrait only), remove NSAllowsLocalNetworking', '1h'],
            ['3',   'Auth — Apple',          'Install package, add capability, create backend endpoint, implement component, test on device', '6–8h'],
            ['4',   'Auth — Google',         'Complete BUG-019 fix, move client ID to env, add URL scheme, test end-to-end', '3–4h'],
            ['4',   'Security',              'BUG-008 (JWT), BUG-011 (SSRF), BUG-016 (admin endpoint), rate limiting', '6–8h'],
            ['5',   'TestFlight Build',      'Archive, upload, wait for processing, invite internal testers, fix build errors', '4–6h'],
            ['6',   'QA & Bug Fixes',        'BUG-012 (NetInfo), BUG-013 (QueryClient), BUG-015 (placeholders), BUG-017 (encryption)', '4–6h'],
            ['7',   'App Store Prep',        'Screenshots, description, privacy policy, age rating, metadata, review notes', '4–6h'],
            ['8',   'Review & Submit',       'Final TestFlight test, submit for App Store review, monitor review status', '2–3h'],
            ['10–12','App Store Review',     'Apple review period (typically 1–3 business days)', '0h (waiting)'],
            ['12+', 'Live on App Store',     'Monitor Crashlytics, Firebase Analytics, and Render logs for launch issues', 'Ongoing'],
        ],
        col_widths=[0.5, 1.7, 4.0, 1.0],
    )
    doc.add_page_break()


def section_app_store_metadata(doc):
    heading(doc, '7  App Store Metadata Requirements', level=1, color=C_DARK_GREY)
    body(doc, 'All fields must be completed in App Store Connect before submission. Items marked Required will block submission if missing.', size=10)

    make_table(doc,
        ['Field', 'Required', 'Value / Notes'],
        [
            ['App Name',           'Yes', 'Shazlo (max 30 chars)'],
            ['Subtitle',           'No',  'Swipe to Discover Fashion (max 30 chars)'],
            ['Bundle ID',          'Yes', 'com.shazlo.mobile'],
            ['SKU',                'Yes', 'SHAZLO-IOS-001 (any unique string)'],
            ['Category (Primary)', 'Yes', 'Shopping'],
            ['Category (Secondary)','No', 'Lifestyle'],
            ['Version',            'Yes', '1.0'],
            ['Description',        'Yes', 'Min 4 sentences. Explain the swipe-to-discover mechanism, AI recommendations, and fashion categories.'],
            ['Keywords',           'Yes', 'fashion,shopping,style,swipe,clothing,outfits,discover,wardrobe,AI (max 100 chars total)'],
            ['Support URL',        'Yes', 'https://www.shazlo.store/support'],
            ['Privacy Policy URL', 'Yes', 'https://www.shazlo.store/privacy (must be live before submission)'],
            ['Marketing URL',      'No',  'https://www.shazlo.store'],
            ['Screenshots 6.9"',   'Yes', '1320×2868px — minimum 3, maximum 10'],
            ['Screenshots 6.5"',   'Yes', '1242×2688px — minimum 3, maximum 10'],
            ['App Preview Video',  'No',  'MP4, max 30s, 1920×1080'],
            ['Age Rating',         'Yes', 'Answer questionnaire — expected result: 4+'],
            ['Copyright',          'Yes', '© 2026 Shazlo Technologies'],
            ['Review Notes',       'Rec.','Explain: shadow user flow, demo account if needed, Razorpay test mode'],
            ['Demo Account',       'Rec.','Provide test email/password if any screen requires login to reach'],
        ],
        col_widths=[1.8, 0.9, 4.1],
    )
    doc.add_page_break()


def section_privacy(doc):
    heading(doc, '8  Privacy & Legal Compliance', level=1, color=C_DARK_GREY)
    body(doc, 'Apple requires accurate disclosure of all data collected. Mismatches between PrivacyInfo.xcprivacy and actual behaviour will cause rejection.', size=10)

    heading(doc, '8.1  Data Collected vs Disclosed', level=2)
    make_table(doc,
        ['Data Type', 'Currently Collected', 'Disclosed in PrivacyInfo', 'Fix Required'],
        [
            ['Name',               'Yes (signup)',          'No',   'Add NSPrivacyCollectedDataTypes entry'],
            ['Email address',       'Yes (signup)',         'No',   'Add NSPrivacyCollectedDataTypes entry'],
            ['Phone number',        'Yes (signup)',         'No',   'Add NSPrivacyCollectedDataTypes entry'],
            ['Date of birth',       'Yes (profile)',        'No',   'Add NSPrivacyCollectedDataTypes entry'],
            ['Gender',              'Yes (profile)',        'No',   'Add NSPrivacyCollectedDataTypes entry'],
            ['Swipe behavior',      'Yes (all swipes)',     'No',   'Add NSPrivacyCollectedDataTypes entry'],
            ['Purchase history',    'Yes (cart/orders)',    'No',   'Add NSPrivacyCollectedDataTypes entry'],
            ['FCM push token',      'Yes (notifications)',  'No',   'Add NSPrivacyCollectedDataTypes entry'],
            ['College/Student',     'Yes (profile)',        'No',   'Add NSPrivacyCollectedDataTypes entry'],
            ['Device identifiers',  'No (no IDFA)',         'N/A',  'None'],
            ['Location',            'Requested only',       'No',   'Add if permission is used'],
        ],
        col_widths=[1.8, 1.7, 1.7, 2.0],
    )

    heading(doc, '8.2  Required Privacy Policy Sections', level=2)
    items = [
        'What data is collected (name, email, phone, swipe behavior, purchase history)',
        'How data is used (personalised recommendations, order fulfilment, marketing)',
        'Third parties data is shared with (Firebase/Google Analytics, Razorpay, Cloudinary)',
        'How users can delete their account and all associated data',
        'Data retention period',
        'Contact email for privacy inquiries',
        'GDPR rights (if EU users): right to access, rectify, erase, portability',
        'CCPA rights (if California users): right to know, delete, opt-out of sale',
        'How to opt out of marketing communications',
        'Cookie/tracking policy',
    ]
    for item in items:
        bullet(doc, item, size=9.5)
    doc.add_page_break()


def section_common_errors(doc):
    heading(doc, '9  Common iOS Build Errors & Fixes', level=1, color=C_DARK_GREY)
    body(doc, 'Reference guide for errors you will encounter during the iOS build process.', size=10)

    errors = [
        {
            'error': "error: The sandbox is not in sync with the Podfile.lock",
            'cause': 'Podfile.lock is out of date with installed pods',
            'fix': 'cd app/ios && pod install --repo-update',
        },
        {
            'error': "FirebaseApp has not been configured. Please call FirebaseApp.configure()",
            'cause': 'BUG-003 — FirebaseApp.configure() not called in AppDelegate',
            'fix': 'Apply the full AppDelegate.swift replacement from BUG-003',
        },
        {
            'error': "GoogleService-Info.plist file not found",
            'cause': 'File not added to Xcode target',
            'fix': 'Drag file into Xcode, check "Copy items if needed", verify in Copy Bundle Resources',
        },
        {
            'error': "Invalid hook call. Hooks can only be called inside of the body of a function component",
            'cause': 'BUG-006 — useSelector inside useEffect',
            'fix': 'Move all useSelector calls to component top level',
        },
        {
            'error': "App Transport Security has blocked a cleartext HTTP resource load",
            'cause': 'API_BASE_URL still uses http:// instead of https://',
            'fix': 'Apply BUG-001 fix — change to https://api.shazlo.store',
        },
        {
            'error': "No account for team. Add a new account in Xcode → Preferences → Accounts",
            'cause': 'Apple Developer account not added to Xcode',
            'fix': 'Xcode → Settings → Accounts → + → Apple ID → sign in with developer account',
        },
        {
            'error': "Provisioning profile doesn\'t include the currently selected device",
            'cause': 'Test device not registered in Apple Developer Portal',
            'fix': 'Go to developer.apple.com → Devices → Register Device → add UDID of test iPhone',
        },
        {
            'error': "Build input files cannot be found: GoogleService-Info.plist",
            'cause': 'File in filesystem but not in Xcode project',
            'fix': 'Delete reference from Xcode and re-add: right-click Shazlo group → Add Files to "Shazlo"',
        },
        {
            'error': "ld: library not found for -lFlipperKit",
            'cause': 'Flipper reference left in Podfile after removal',
            'fix': 'Check Podfile for any Flipper references and remove them. Run pod install',
        },
        {
            'error': "RCTBridge required dispatch_sync to load",
            'cause': 'Old Bridge architecture mixed with New Architecture (Fabric)',
            'fix': 'Ensure RCTNewArchEnabled=1 in Podfile post_install and all native modules support New Architecture',
        },
    ]

    make_table(doc,
        ['Error', 'Cause', 'Fix'],
        [[e['error'], e['cause'], e['fix']] for e in errors],
        col_widths=[2.0, 1.8, 2.9],
    )
    doc.add_page_break()


def section_final_checklist(doc):
    heading(doc, '10  Pre-Submission Final Checklist', level=1, color=C_DARK_GREY)
    body(doc, 'Every item must be checked before clicking "Submit for Review" in App Store Connect.', size=10)

    sections = [
        ('Frontend Fixes', [
            'BUG-001: API_BASE_URL uses https://api.shazlo.store (not local IP)',
            'BUG-006: useSelector is at component top level (not inside useEffect)',
            'BUG-012: Single NetInfo listener with cleanup',
            'BUG-013: QueryClient declared at module level',
            'BUG-014: No PII in URL query parameters (Payment screen)',
            'BUG-015: No hardcoded placeholder product details in SwipeUI',
            'BUG-018: console.log stripped from production builds (babel plugin)',
            'BUG-019: Google Sign-In wired to backend and Redux',
            'console.log(items) removed from SwipeUI.jsx',
        ]),
        ('iOS Native', [
            'GoogleService-Info.plist added to Xcode project and Copy Bundle Resources',
            'AppDelegate.swift includes FirebaseApp.configure() as first call',
            'NSLocationWhenInUseUsageDescription has meaningful text (not empty)',
            'NSAllowsLocalNetworking removed from Info.plist',
            'Only UIInterfaceOrientationPortrait in UISupportedInterfaceOrientations',
            'Sign In with Apple capability enabled in Xcode',
            'Push Notifications capability enabled in Xcode',
            'Associated Domains capability: applinks:www.shazlo.store',
            'APNs key uploaded to Firebase Console',
            'REVERSED_CLIENT_ID URL scheme added to Info.plist for Google Sign-In',
            'pod install ran successfully with no errors',
            'App builds and runs on real iPhone without crash',
        ]),
        ('Authentication', [
            'Sign in with Apple implemented and tested on real device',
            'Google Sign-In completes end-to-end (SDK → backend → Redux)',
            'Shadow user flow still works (continue without login)',
            'JWT tokens issued and stored after all auth flows',
        ]),
        ('Backend Production', [
            'DEBUG=False deployed to Render',
            'ALLOWED_HOSTS=api.shazlo.store deployed',
            'CORS_ALLOW_ALL_ORIGINS=False deployed',
            'Print statements with credentials removed',
            'BUG-009: email bug fixed and deployed',
            'Rate limiting on shadow user and login endpoints',
            'SSRF proxy has domain allowlist',
            'Admin endpoint protected with IsAdminUser',
        ]),
        ('App Store Connect', [
            'App name, subtitle, description, keywords filled in',
            'Privacy Policy URL is live and accessible',
            'Screenshots uploaded (6.9" and 6.5" sizes)',
            'Age rating questionnaire completed',
            'Review notes explain shadow user flow',
            'NSPrivacyCollectedDataTypes in PrivacyInfo.xcprivacy matches actual data collected',
            'Version set to 1.0, Build number incremented for each upload',
        ]),
    ]

    for section_name, items in sections:
        heading(doc, section_name, level=2, color=C_DARK_GREY)
        for item in items:
            bullet(doc, f'☐  {item}', size=9.5)
        doc.add_paragraph()

    doc.add_page_break()


def footer_page(doc):
    heading(doc, 'Document Information', level=1, color=C_DARK_GREY)
    make_table(doc,
        ['Field', 'Value'],
        [
            ['Document Title',    'Shazlo iOS App Store Production Engineering Audit'],
            ['Version',           '1.0'],
            ['Generated',         str(date.today())],
            ['App',               'Shazlo — AI Fashion Discovery (React Native 0.80.0)'],
            ['Backend',           'Django 5.2.3 — https://api.shazlo.store/'],
            ['Bundle ID',         'com.shazlo.mobile'],
            ['Total Bugs Found',  '20 (8 Critical, 6 High, 6 Medium)'],
            ['Estimated Fix Time','8–12 engineering days'],
            ['Verdict',           'NOT READY FOR APP STORE — resolve all critical bugs first'],
            ['Next Review',       'After all critical + high bugs are resolved'],
        ],
        col_widths=[2.0, 5.3],
    )
    body(doc, '')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '— End of Document —', bold=True, size=10, color=C_GOLD)


# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page margins — 1 inch all sides
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    cover_page(doc)
    section_executive_summary(doc)
    section_bug_registry(doc)
    section_ios_deployment(doc)
    section_backend_hardening(doc)
    section_costing(doc)
    section_timeline(doc)
    section_app_store_metadata(doc)
    section_privacy(doc)
    section_common_errors(doc)
    section_final_checklist(doc)
    footer_page(doc)

    out_path = os.path.join(os.path.dirname(__file__), 'Shazlo_Engineering_Audit.docx')
    doc.save(out_path)
    print(f'✅  Document saved: {out_path}')
    print(f'    Pages (approx): {len(doc.paragraphs) // 40 + 1}')
    print('    Open in: Microsoft Word, Google Docs (File → Import), LibreOffice')


if __name__ == '__main__':
    main()
